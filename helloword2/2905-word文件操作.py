import docx
import random
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

doc = docx.Document()
doc.styles['Normal'].font.name = '华文行楷'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '华文行楷')

for i in range(3):
    par = doc.add_paragraph()
    for _ in range(5):
        run = par.add_run('文本颜色测试(test)')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(5 + 2 * i)

        run.font.bold = True
        run.font.italic = True
        run.font.underline = False

        color = (random.randint(0, 255) for _ in range(3))
        run.font.color.rgb = RGBColor(*color)
doc.save('2905.docx')
