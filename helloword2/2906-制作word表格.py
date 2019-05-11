import docx
import random
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

wb = openpyxl.load_workbook('工资处理.xlsx')
sheet = wb['工作量汇总']
lst = []
for i, r in enumerate(sheet.rows, start=0):
    if i < 13:
        lst.append([cell.value for cell in r])
doc = docx.Document()
doc.styles['Normal'].font.name = '宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')



par = doc.add_paragraph()
run = par.add_run('工资表')
run.font.size = Pt(30)
run.font.bold = True
par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


table = doc.add_table(rows=len(lst), cols=len(lst[0]), style='Table Grid')
for r in range(len(lst)):
    for c in range(len(lst[0])):
        if r > 0 and c > 1:
            lst[r][c] = int(lst[r][c] * 100 + 0.5) / 100
        cell = table.cell(r, c)
        cell.text = str(lst[r][c])
        par = cell.paragraphs[0]
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = par.runs[0]
        run.font.name = 'Times New Roman'
        run.font.bold = False
        run.font.italic = False
        run.font.underline = False
        color = (random.randint(0, 255) for _ in range(3))
        run.font.color.rgb = RGBColor(*color)
doc.save('2906工作量汇总.docx')
